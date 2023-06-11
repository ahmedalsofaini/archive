from docx import Document
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm
import os

# Constants for page and font properties
PAGE_HEIGHT = Mm(470)
PAGE_WIDTH = Mm(210)
LEFT_MARGIN = Mm(18)
RIGHT_MARGIN = Mm(22)
TOP_MARGIN = Mm(12.7)
BOTTOM_MARGIN = Mm(20)
FONT_NAME = 'Sahel'
FONT_SIZE = Pt(16)

def doc_fun():
    """Create a new Word document and set its page and font properties.

    Returns:
        document: a Document object.
    """
    document = Document()

    section = document.sections[0]
    section.page_height = PAGE_HEIGHT
    section.page_width = PAGE_WIDTH
    section.left_margin = LEFT_MARGIN
    section.right_margin = RIGHT_MARGIN
    section.top_margin = TOP_MARGIN
    section.bottom_margin = BOTTOM_MARGIN
    section.alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT # Align the section to the right
    font = document.styles['Normal'].font # Get the font of the normal style
    font.name = FONT_NAME # Set the font name to Sahel
    font.size = FONT_SIZE # Set the font size to 16

    return document

def fetch(pathf,full_text,png_path_and_name):
    try:
        # Create the document object by calling the doc_fun function
        document=doc_fun()
        
        # Add a paragraph with the given text and align it to the right
        p1 = document.add_paragraph(full_text)
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add a table with 2 rows and 5 columns and align it to the right
        table = document.add_table(rows=2, cols=5) # Create a table with two rows and five columns
        table.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT # Align the table to the right

        # Access the cells of the table by indexing
        cell_1 = table.cell(0, 0) # First cell of the first row
        cell_2 = table.cell(0, 1) # Second cell of the first row
        cell_3 = table.cell(0, 2) # Third cell of the first row
        cell_4 = table.cell(0, 3) # Fourth cell of the first row
        cell_5 = table.cell(0, 4) # Fifth cell of the first row
        
        cell_6 = table.cell(1, 0) # First cell of the second row
        cell_7 = table.cell(1, 1) # Second cell of the second row
        cell_8 = table.cell(1, 2) # Third cell of the second row
        cell_9 = table.cell(1, 3) # Fourth cell of the second row
        cell_10 = table.cell(1, 4) # Fifth cell of the second row

        # Add the content to each cell using paragraphs and runs

        p2 = cell_1.add_paragraph('السنة') # Add the text 'السنة' to the first cell of the first row
        
        p3 = cell_2.add_paragraph('الدفتر') # Add the text 'الدفتر' to the second cell of the first row
        
        p4 = cell_3.add_paragraph('الصفحة') # Add the text 'الصفحة' to the third cell of the first row
        
        p5 = cell_4.add_paragraph('الملف') # Add the text 'الملف' to the fourth cell of the first row
        
        p6 = cell_5.add_paragraph('الباركود') # Add the text 'الباركود' to the fifth cell of the first row
        
        p7 = cell_6.add_paragraph('1444') # Add the text '1444' to the first cell of the second row
        
        p8 = cell_7.add_paragraph('20') # Add the text '20' to the second cell of the second row
        
        p9 = cell_8.add_paragraph('40') # Add the text '40' to the third cell of the second row
        
        p10 = cell_9.add_paragraph('222') # Add the text '222' to the fourth cell of the second row
        
        p11 = cell_10.add_paragraph() # Add a new paragraph to the fifth cell of the second row
        r1 = p11.add_run() # Add a new run to the paragraph
        r1.add_picture(png_path_and_name, width=Inches(0.5)) # Add the picture to the run with half size

        table.style = 'Table Grid' # Set the table style to Table Grid
        table.autofit = True # Set the table autofit to True

        # Remove the space before each paragraph in the table cells
        for cell in table._cells: # Loop through each cell in the table
            for paragraph in cell.paragraphs: # Loop through each paragraph in the cell
                paragraph.paragraph_format.space_before = Pt(0) # Set the space before to zero

        # Save the document in the given path
        document.save(pathf)
        
        # Print the directory name where the document is saved
        print('Directory Name:     ',pathf )
    
    except Exception as e:
        # Print the error message if any exception occurs
        print('An error occurred:', e)
